import logging
import os
import re
from typing import Dict, Any, List, Optional

try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    logging.error("python-docx not installed. Install it with 'pip install python-docx'")

from app.core.parser_interface import FileParserInterface, ParserResult

logger = logging.getLogger(__name__)


class DOCXParser(FileParserInterface):
    """
    Parser for Microsoft Word (.docx) files.
    """
    
    def __init__(self):
        """Initialize the DOCX parser."""
        # Pattern to identify potential section headings from style names
        self.heading_styles = ['Heading', 'Title', 'Subtitle']
    
    def parse(self, file_path: str) -> ParserResult:
        """
        Parse a DOCX file and return its structured content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A ParserResult object containing the parsed content
        """
        logger.info(f"Parsing DOCX file: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
            
        try:
            # Ensure python-docx is installed
            if 'docx' not in globals():
                logger.error("python-docx library not available. Install with 'pip install python-docx'")
                raise ImportError("python-docx library not available")
                
            text = self.get_full_text(file_path)
            metadata = self._extract_metadata(file_path)
            sections = self.get_sections(file_path)
            
            return ParserResult(
                text=text,
                sections=sections,
                metadata=metadata,
                file_path=file_path
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise
    
    def get_full_text(self, file_path: str) -> str:
        """
        Extract the full text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            The extracted text content
        """
        try:
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return full_text
            
        except PackageNotFoundError:
            logger.error(f"Invalid DOCX file or file corrupted: {file_path}")
            raise
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file {file_path}: {str(e)}")
            raise
    
    def get_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract sections from a DOCX file based on heading styles.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A list of sections, each containing title and content
        """
        try:
            doc = docx.Document(file_path)
            sections = []
            current_section: Optional[Dict[str, Any]] = None
            current_content = []
            
            # Process each paragraph
            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if self._is_heading(para):
                    # If we already have a section, save it
                    if current_section:
                        current_section["content"] = "\n".join(current_content)
                        sections.append(current_section)
                    
                    # Start a new section
                    current_section = {
                        "title": para.text.strip(),
                        "content": ""
                    }
                    current_content = []
                elif current_section is not None:
                    # Add to current section content if it's not empty
                    if para.text.strip():
                        current_content.append(para.text)
                elif para.text.strip():
                    # Handle content before any headings
                    # Create an introduction section
                    current_section = {
                        "title": "Introduction",
                        "content": ""
                    }
                    current_content = [para.text]
            
            # Add the last section
            if current_section:
                current_section["content"] = "\n".join(current_content)
                sections.append(current_section)
                
            # If no sections were found, create a single section with all content
            if not sections:
                full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                sections = [{"title": "Content", "content": full_text}]
                
            return sections
            
        except Exception as e:
            logger.error(f"Error extracting sections from DOCX file {file_path}: {str(e)}")
            raise
    
    def _extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A dictionary containing metadata
        """
        try:
            doc = docx.Document(file_path)
            
            metadata = {
                "file_size": os.path.getsize(file_path),
                "page_count": self._estimate_page_count(doc),
            }
            
            # Extract core properties
            core_props = doc.core_properties
            if core_props:
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["created"] = core_props.created.isoformat() if hasattr(core_props.created, 'isoformat') else str(core_props.created)
                if core_props.modified:
                    metadata["modified"] = core_props.modified.isoformat() if hasattr(core_props.modified, 'isoformat') else str(core_props.modified)
                if core_props.comments:
                    metadata["comments"] = core_props.comments
                    
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting metadata from DOCX file {file_path}: {str(e)}")
            # Return basic file metadata if document metadata extraction fails
            return {
                "file_size": os.path.getsize(file_path)
            }
    
    def _is_heading(self, paragraph) -> bool:
        """
        Determine if a paragraph is a heading based on its style.
        
        Args:
            paragraph: A docx paragraph object
            
        Returns:
            True if the paragraph is a heading, False otherwise
        """
        if not paragraph.text.strip():
            return False
            
        # Check style name
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            for heading_style in self.heading_styles:
                if heading_style.lower() in style_name:
                    return True
            
            # Check if it's a numbered heading style (Heading 1, Heading 2, etc.)
            if re.match(r'heading \d+', style_name):
                return True
                
        # Heuristic: If text is short and followed by a longer paragraph,
        # it might be a heading
        if len(paragraph.text) < 100 and paragraph.text.strip().endswith((':','.')) is False:
            # Additional checks could be implemented here
            pass
            
        return False
    
    def _estimate_page_count(self, doc: DocxDocument) -> int:
        """
        Estimate the page count of a DOCX document.
        This is an approximation since actual page count depends on formatting.
        
        Args:
            doc: A docx Document object
            
        Returns:
            Estimated page count
        """
        # Rough approximation: about 500 words per page
        word_count = sum(len(para.text.split()) for para in doc.paragraphs)
        page_count = max(1, round(word_count / 500))
        
        return page_count 